import os
from docx import Document
import pdfplumber
import re
from typing import Dict

def extract_text_from_pdf(file) -> str:
    with pdfplumber.open(file) as pdf:
        text = ''
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + '\n'
    return text

def extract_text_from_docx(file) -> str:
    doc = Document(file)
    text = '\n'.join([para.text for para in doc.paragraphs])
    return text

def extract_resume_info(text: str) -> Dict[str, str]:
    # 简单正则提取，实际可根据简历模板优化
    name = re.search(r'Name[:：]?\s*([A-Za-z\u4e00-\u9fa5\s]+)', text)
    email = re.search(r'[\w\.-]+@[\w\.-]+', text)
    major = re.search(r'Major[:：]?\s*([A-Za-z\u4e00-\u9fa5\s]+)', text)
    skills = re.search(r'Skills?[:：]?\s*([A-Za-z,\u4e00-\u9fa5,\s]+)', text)
    return {
        'name': name.group(1).strip() if name else '',
        'email': email.group(0).strip() if email else '',
        'major': major.group(1).strip() if major else '',
        'skill': skills.group(1).strip() if skills else ''
    }

def parse_resume(file, filename):
    """
    解析简历文件，支持 PDF 和 DOCX 格式
    Args:
        file: FileStorage 对象或文件路径
        filename: 文件名，用于判断文件类型
    Returns:
        dict: 包含解析出的信息
        {
            'name': '姓名',
            'email': '邮箱',
            'major': '专业',
            'skill': '技能'
        }
    """
    # 获取文件扩展名
    ext = os.path.splitext(filename)[1].lower()
    
    # 根据文件类型选择不同的解析方法
    if ext == '.pdf':
        return parse_pdf(file)
    elif ext in ['.docx', '.doc']:
        return parse_docx(file)
    else:
        raise ValueError('不支持的文件格式，仅支持 PDF 和 DOCX 格式')

def parse_pdf(file):
    """解析 PDF 文件"""
    text = ""
    with pdfplumber.open(file) as pdf:
        for page in pdf.pages:
            text += page.extract_text() or ""
    
    return extract_info_from_text(text)

def parse_docx(file):
    """解析 DOCX 文件"""
    doc = Document(file)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    
    return extract_info_from_text(text)

def extract_info_from_text(text):
    """
    从英文文本中提取关键信息
    优先用关键词匹配，否则用正则和启发式推断
    """
    lines = text.split('\n')
    info = {
        'name': '',
        'email': '',
        'major': '',
        'skill': ''
    }
    import re
    email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\\.[a-zA-Z]{2,}'
    skill_keywords = ['python', 'java', 'c++', 'c#', 'sql', 'html', 'css', 'js', 'javascript', 'machine learning', 'deep learning', 'data analysis', 'frontend', 'backend', 'linux', 'git', 'docker', 'cloud', 'tensorflow', 'pytorch']
    major_keywords = ['computer science', 'software', 'engineering', 'information', 'automation', 'electrical', 'mathematics', 'physics', 'ai', 'artificial intelligence', 'data science', 'cybersecurity', 'network', 'informatics']

    # 全局找邮箱
    email_match = re.search(email_pattern, text, re.IGNORECASE)
    if email_match:
        info['email'] = email_match.group()

    # 逐行找Name、Major、Skill关键词，否则用启发式
    for idx, line in enumerate(lines):
        l = line.strip()
        if not l:
            continue
        # Name: xxx
        if not info['name']:
            m = re.match(r'Name[:：]?\s*([A-Za-z .-]+)', l)
            if m:
                info['name'] = m.group(1).strip()
        # Major: xxx
        if not info['major']:
            m = re.match(r'Major[:：]?\s*([A-Za-z .-]+)', l)
            if m:
                info['major'] = m.group(1).strip()
        # Skills: xxx
        if not info['skill']:
            m = re.match(r'Skills?[:：]?\s*([A-Za-z0-9,./+\- ]+)', l, re.IGNORECASE)
            if m:
                info['skill'] = m.group(1).strip()

    # 启发式补全英文名（首行大写单词，且不是邮箱/技能/专业）
    if not info['name']:
        for l in lines:
            l = l.strip()
            if l and not re.search(email_pattern, l) and len(l.split()) <= 4:
                # 只包含字母和空格，且首字母大写
                if re.match(r'^[A-Z][a-zA-Z .-]+$', l):
                    info['name'] = l
                    break
    # 启发式补全专业
    if not info['major']:
        for l in lines:
            for kw in major_keywords:
                if kw in l.lower():
                    info['major'] = kw
                    break
            if info['major']:
                break
    # 启发式补全技能
    if not info['skill']:
        found_skills = []
        for l in lines:
            for kw in skill_keywords:
                if kw in l.lower() and kw not in found_skills:
                    found_skills.append(kw)
        if found_skills:
            info['skill'] = ','.join(found_skills)
    return info 